"""Валидация .docx, сгенерированного gost_report, на соответствие ГОСТ 7.32.

Два слоя потребляют этот модуль:
  L1: Report.save() импортирует validate_docx() и raise'ит GostValidationError
      на failures — модель видит traceback и сама исправляет, без какой-либо
      инструкции про валидацию в SKILL.md.
  L2: Stop-hook вызывает `python3 validate.py --hook`. Hook сканирует cwd на
      sentinel-файлы (.gost-meta.json), валидирует соответствующие .docx,
      возвращает {"decision": "block", "reason": "..."} JSON если есть
      нарушения. Hook никогда не валит Stop pipeline — на любой ошибке exit 0.

Sentinel format: <docx_path>.gost-meta.json — пишется Report.save() рядом с
docx, нужен чтобы L2-хук срабатывал ТОЛЬКО на gost-report-генерации, а не
на любой .docx в cwd (false-fires в чужих проектах).
"""
from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
import traceback
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    Document = None
    qn = None
    HAS_DOCX = False


def _venv_python() -> Optional[Path]:
    """Путь к скилл-venv питону, если он есть рядом."""
    skill_dir = Path(__file__).resolve().parent.parent
    if os.name == "nt":
        candidate = skill_dir / ".venv" / "Scripts" / "python.exe"
    else:
        candidate = skill_dir / ".venv" / "bin" / "python"
    return candidate if candidate.exists() else None


def _maybe_reexec_in_venv(argv: List[str]) -> None:
    """Если python-docx не доступен в текущем питоне, но скилл-venv существует —
    re-exec под venv-питоном. Используется только в CLI-режиме (--hook / --check):
    при `import validate` из gost_report этот код не выполняется (importer уже
    на правильном питоне).

    Если venv нет (юзер ещё ни разу не вызывал r.save()) — функция возвращает
    управление; CLI решит сам, что делать (--hook silently exit 0, --check
    прокинет missing-deps Violation).
    """
    if HAS_DOCX:
        return
    venv_py = _venv_python()
    if venv_py is None:
        return
    try:
        same = Path(sys.executable).resolve() == venv_py.resolve()
    except OSError:
        same = False
    if same:
        return
    script = str(Path(__file__).resolve())
    if os.name == "nt":
        rc = subprocess.run([str(venv_py), script, *argv]).returncode
        sys.exit(rc)
    os.execv(str(venv_py), [str(venv_py), script, *argv])


SENTINEL_VERSION = 1
SENTINEL_SUFFIX = ".gost-meta.json"
MAX_VIOLATIONS_IN_REPORT = 10
HOOK_GLOB_MAX_DEPTH = 6

TIER_FAIL = "a"
TIER_WARN = "b"


@dataclass(frozen=True)
class Violation:
    tier: str        # "a" hard fail, "b" heuristic warn
    code: str        # short rule id
    message: str     # ru, one sentence
    location: str = ""

    def format(self) -> str:
        prefix = "[FAIL]" if self.tier == TIER_FAIL else "[WARN]"
        loc = f" ({self.location})" if self.location else ""
        return f"  {prefix} [{self.code}]{loc}: {self.message}"


# ---- Sentinel ---------------------------------------------------------------

def _sentinel_path(docx_path: Path) -> Path:
    return docx_path.parent / (docx_path.name + SENTINEL_SUFFIX)


def write_sentinel(docx_path: Path, *, profile_name: str = "",
                   generator_version: str = "") -> None:
    payload = {
        "version": SENTINEL_VERSION,
        "generator": "gost_report",
        "generator_version": generator_version,
        "profile": profile_name,
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "docx": str(docx_path),
    }
    sentinel = _sentinel_path(docx_path)
    sentinel.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def read_sentinel(docx_path: Path) -> Optional[dict]:
    sentinel = _sentinel_path(docx_path)
    if not sentinel.exists():
        return None
    try:
        return json.loads(sentinel.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return None


# ---- Paragraph classification ----------------------------------------------

_FIGURE_CAPTION_RE = re.compile(r"^Рисунок (\d+) — ")
_TABLE_CAPTION_RE = re.compile(r"^Таблица (\d+) — ")


def _para_is_caption(text: str) -> bool:
    return bool(_FIGURE_CAPTION_RE.match(text)
                or _TABLE_CAPTION_RE.match(text))


def _para_is_code(paragraph) -> bool:
    for run in paragraph.runs:
        if run.font.name == "Courier New":
            return True
    return False


def _para_has_drawing(paragraph) -> bool:
    if qn is None:
        return False
    for _ in paragraph._element.iter(qn("w:drawing")):
        return True
    return False


def _para_text_from_xml(p_element) -> str:
    if qn is None:
        return ""
    parts = []
    for t in p_element.iter(qn("w:t")):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


# ---- Tier (a) checks: hard fail --------------------------------------------

_PROHIBITED_DASH_RE = re.compile(r"[—–]")


def _check_dashes(doc) -> List[Violation]:
    violations: List[Violation] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        if not text:
            continue
        if _para_is_caption(text):
            continue
        if _para_is_code(p):
            continue
        m = _PROHIBITED_DASH_RE.search(text)
        if m:
            start = max(0, m.start() - 20)
            end = min(len(text), m.end() + 20)
            snippet = text[start:end]
            violations.append(Violation(
                tier=TIER_FAIL, code="dashes",
                message=f"длинное/среднее тире в основном тексте: «...{snippet}...»",
                location=f"параграф {i + 1}",
            ))
    return violations


_BARE_LATEX_RE = re.compile(
    r"\\(?:frac|sum|int|prod|sqrt|text|begin|end|cdot|times|div|"
    r"alpha|beta|gamma|delta|epsilon|zeta|eta|theta|iota|kappa|lambda|"
    r"mu|nu|xi|omicron|pi|rho|sigma|tau|upsilon|phi|chi|psi|omega|"
    r"infty|partial|nabla|hat|bar|tilde|vec|dot|left|right|mathbb|mathcal)"
    r"\b"
)


def _check_bare_latex(doc) -> List[Violation]:
    violations: List[Violation] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        if not text or _para_is_code(p):
            continue
        m = _BARE_LATEX_RE.search(text)
        if m:
            violations.append(Violation(
                tier=TIER_FAIL, code="bare-latex",
                message=(f"LaTeX-команда {m.group(0)!r} попала в текст. "
                         "Используй r.formula(...) для уравнений."),
                location=f"параграф {i + 1}",
            ))
    return violations


def _check_figure_captions(doc) -> List[Violation]:
    violations: List[Violation] = []
    paragraphs = doc.paragraphs
    expected = 1
    for i, p in enumerate(paragraphs):
        if not _para_has_drawing(p):
            continue
        # Caption — следующий непустой параграф (макс 2 шага вперёд).
        caption = ""
        for j in range(i + 1, min(i + 3, len(paragraphs))):
            if paragraphs[j].text.strip():
                caption = paragraphs[j].text
                break
        m = _FIGURE_CAPTION_RE.match(caption)
        if not m:
            violations.append(Violation(
                tier=TIER_FAIL, code="figure-no-caption",
                message=("у рисунка нет подписи «Рисунок N — Описание» "
                         "сразу после изображения."),
                location=f"параграф {i + 1}",
            ))
            continue
        n = int(m.group(1))
        if n != expected:
            violations.append(Violation(
                tier=TIER_FAIL, code="figure-numbering",
                message=(f"номер рисунка {n}, ожидался {expected}. "
                         "Нумерация не должна прыгать."),
                location=f"параграф {i + 1}",
            ))
        expected = n + 1
    return violations


def _check_table_captions(doc) -> List[Violation]:
    """Каждая таблица должна иметь подпись «Таблица N — ...» прямо перед ней.

    Плавающие borderless-таблицы (городкод/год на титульнике, у них
    `<w:tblpPr>` внутри `<w:tblPr>`) пропускаются — у них нет подписи по
    дизайну.
    """
    violations: List[Violation] = []
    if qn is None:
        return violations
    body = doc.element.body
    expected = 1
    last_nonempty_text = ""
    for child in body:
        local = child.tag.split("}", 1)[-1]
        if local == "p":
            text = _para_text_from_xml(child)
            if text.strip():
                last_nonempty_text = text
        elif local == "tbl":
            tblpr = child.find(qn("w:tblPr"))
            is_floating = (tblpr is not None
                           and tblpr.find(qn("w:tblpPr")) is not None)
            if is_floating:
                last_nonempty_text = ""
                continue
            m = _TABLE_CAPTION_RE.match(last_nonempty_text)
            if not m:
                violations.append(Violation(
                    tier=TIER_FAIL, code="table-no-caption",
                    message=("у таблицы нет подписи «Таблица N — Описание» "
                             "прямо перед ней."),
                ))
            else:
                n = int(m.group(1))
                if n != expected:
                    violations.append(Violation(
                        tier=TIER_FAIL, code="table-numbering",
                        message=f"номер таблицы {n}, ожидался {expected}.",
                    ))
                expected = n + 1
            last_nonempty_text = ""
    return violations


_PLACEHOLDER_NAME_RE = re.compile(r"Фамилия\s+И\.О\.")


def _check_placeholder_name(doc) -> List[Violation]:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if _PLACEHOLDER_NAME_RE.search(full_text):
        return [Violation(
            tier=TIER_FAIL, code="placeholder-name",
            message=("«Фамилия И.О.» из примеров skill'а попало в документ. "
                     "Подставь реальное ФИО студента/преподавателя."),
        )]
    return []


# ---- Tier (b) checks: heuristic warn ---------------------------------------

_AI_TONE_PATTERNS = [
    re.compile(r"в\s+ходе\s+выполнения\s+работ\w*", re.IGNORECASE),
    re.compile(r"в\s+результате\s+проведённого\s+исследовани\w*", re.IGNORECASE),
    re.compile(r"\b(?:данн\w+|вышеуказанн\w+|нижеследующ\w+)\b", re.IGNORECASE),
    re.compile(r"\bпредставляет\s+собой\b", re.IGNORECASE),
    re.compile(r"\bосуществляется\b", re.IGNORECASE),
    re.compile(r"стоит\s+отметить,?\s+что", re.IGNORECASE),
    re.compile(r"необходимо\s+подчеркнуть,?\s+что", re.IGNORECASE),
    re.compile(r"таким\s+образом,?\s+можно\s+сделать\s+вывод", re.IGNORECASE),
]


def _check_ai_tone(doc) -> List[Violation]:
    violations: List[Violation] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        if not text or _para_is_caption(text) or _para_is_code(p):
            continue
        for rx in _AI_TONE_PATTERNS:
            m = rx.search(text)
            if m:
                violations.append(Violation(
                    tier=TIER_WARN, code="ai-tone",
                    message=(f"шаблонная фраза «{m.group(0)}» — "
                             "перепиши естественнее (см. writing-style в SKILL.md)."),
                    location=f"параграф {i + 1}",
                ))
                break
    return violations


def _check_caption_short(doc) -> List[Violation]:
    violations: List[Violation] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        m = _FIGURE_CAPTION_RE.match(text)
        if not m:
            continue
        after = text[m.end():].strip()
        if len(after.split()) < 2:
            violations.append(Violation(
                tier=TIER_WARN, code="caption-short",
                message=(f"подпись рисунка слишком короткая: «{after}». "
                         "Дай осмысленное описание."),
                location=f"параграф {i + 1}",
            ))
    return violations


# ---- Orchestration ---------------------------------------------------------

_ALL_CHECKS = [
    _check_dashes,
    _check_bare_latex,
    _check_figure_captions,
    _check_table_captions,
    _check_placeholder_name,
    _check_ai_tone,
    _check_caption_short,
]


def validate_docx(path: Path) -> List[Violation]:
    if not HAS_DOCX:
        return [Violation(
            tier=TIER_FAIL, code="missing-deps",
            message=("python-docx не установлен; не могу прочитать .docx. "
                     "Запусти через scripts/ensure_env.py."),
        )]
    try:
        doc = Document(str(path))
    except Exception as e:
        return [Violation(
            tier=TIER_FAIL, code="read-failed",
            message=f"не смог открыть .docx: {e}",
        )]
    violations: List[Violation] = []
    for check in _ALL_CHECKS:
        try:
            violations.extend(check(doc))
        except Exception as e:
            violations.append(Violation(
                tier=TIER_WARN, code="rule-error",
                message=f"внутренняя ошибка правила {check.__name__}: {e}",
            ))
    return violations


def has_failures(violations: Iterable[Violation]) -> bool:
    return any(v.tier == TIER_FAIL for v in violations)


def format_error(docx_path: Path, violations: List[Violation]) -> str:
    fails = [v for v in violations if v.tier == TIER_FAIL]
    warns = [v for v in violations if v.tier == TIER_WARN]
    ordered = fails + warns

    lines = [f"Файл {docx_path} не прошёл ГОСТ-проверку:"]
    shown = ordered[:MAX_VIOLATIONS_IN_REPORT]
    for v in shown:
        lines.append(v.format())
    omitted = len(ordered) - len(shown)
    if omitted > 0:
        lines.append(f"  ... и ещё {omitted} "
                     f"(показаны первые {MAX_VIOLATIONS_IN_REPORT}).")
    lines.append("")
    lines.append("Перегенерируй документ: исправь причины и снова вызови r.save().")
    lines.append("Это валидация ГОСТ, не баг библиотеки. Не подавляй это исключение")
    lines.append("и не используй python-docx напрямую — валидация выполняется "
                 "автоматически в r.save().")
    return "\n".join(lines)


# ---- CLI -------------------------------------------------------------------

def _hook_main(cwd: Path) -> int:
    """Stop-хук: сканирует cwd на sentinel'ы, валидирует docx, печатает
    decision-block JSON если есть FAIL'ы. Hook всегда exit 0 — лучше тихо
    промолчать, чем уронить Stop-pipeline."""
    if not HAS_DOCX:
        return 0
    try:
        sentinels = []
        for s in cwd.rglob("*" + SENTINEL_SUFFIX):
            try:
                rel = s.relative_to(cwd)
                if len(rel.parts) <= HOOK_GLOB_MAX_DEPTH:
                    sentinels.append(s)
            except ValueError:
                continue
    except Exception:
        return 0

    if not sentinels:
        return 0

    by_doc: dict = {}
    for sentinel in sentinels:
        try:
            payload = json.loads(sentinel.read_text(encoding="utf-8"))
            docx_str = payload.get("docx") or ""
            docx = Path(docx_str)
            if not docx.is_absolute():
                docx = sentinel.parent / Path(docx_str).name
            if not docx.exists():
                # Sentinel остался от удалённого .docx — игнорируем.
                continue
            violations = validate_docx(docx)
            fails = [v for v in violations if v.tier == TIER_FAIL]
            if fails:
                by_doc[docx] = fails
        except Exception:
            continue

    if not by_doc:
        return 0

    parts = [format_error(d, vs) for d, vs in by_doc.items()]
    reason = "\n\n".join(parts)
    print(json.dumps({"decision": "block", "reason": reason},
                     ensure_ascii=False))
    return 0


def _check_main(path_str: str) -> int:
    """`--check <path>`: ручная проверка одного .docx, человеко-читаемый отчёт.

    Exit: 0 если passed (включая warnings-only), 1 если есть FAIL,
    2 при ошибке вызова.
    """
    path = Path(path_str)
    if not path.exists():
        print(f"validate.py: file not found: {path}", file=sys.stderr)
        return 2
    violations = validate_docx(path)
    if not violations:
        print(f"validate.py: {path} — passed (0 нарушений)")
        return 0
    print(format_error(path, violations))
    return 1 if has_failures(violations) else 0


def main(argv: Optional[List[str]] = None) -> int:
    raw_argv = list(sys.argv[1:]) if argv is None else list(argv)
    _maybe_reexec_in_venv(raw_argv)

    parser = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    parser.add_argument("--hook", action="store_true",
                        help="Stop-hook: сканирует cwd, JSON decision-block")
    parser.add_argument("--check", metavar="PATH",
                        help="ручная проверка одного .docx")
    args = parser.parse_args(argv)

    try:
        if args.hook:
            return _hook_main(Path.cwd())
        if args.check:
            return _check_main(args.check)
        parser.print_help()
        return 2
    except Exception:
        if args.hook:
            return 0
        traceback.print_exc()
        return 2


if __name__ == "__main__":
    sys.exit(main())
