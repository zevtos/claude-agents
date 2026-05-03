"""
gost_report — генератор студенческих/научных работ по ГОСТ 7.32 для любого
российского вуза. Профили университетов задают поля, кегли заголовков,
название университета/факультета и прочие специфичные мелочи; чистый GOST 7.32
работает «из коробки», ИТМО — встроенный пресет.

Минимальный пример (ИТМО, дефолт):
    from gost_report import Report, TitleConfig

    r = Report(TitleConfig(
        work_type="Лабораторная работа",
        work_number="№1",
        topic="Основы работы в командной строке Unix",
        student_name="Фамилия И.О.",
        student_group="P3XXX",
        teacher_name="Фамилия И.О.",
        teacher_degree="к.т.н.",
        teacher_position="доцент",
        year="2026",
    ))
    r.toc()
    r.h1("Введение")
    r.text("Цель работы — изучить...")
    r.save("report.docx")

Любой другой вуз — передайте свой профиль:
    from gost_report import Report, TitleConfig, GOST_PROFILE, UniversityProfile

    SPBSU_PROFILE = UniversityProfile(
        university_short="«Санкт-Петербургский государственный университет»",
        faculty="Математико-механический факультет",
        toc_title="СОДЕРЖАНИЕ",
    )
    r = Report(TitleConfig(...), profile=SPBSU_PROFILE)

Зависимости: pip install python-docx
"""

from dataclasses import dataclass
from pathlib import Path
import re
import sys
import xml.etree.ElementTree as _ET
from typing import List, Optional, Sequence, Union

from docx import Document
from docx.document import Document as _Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _paths import ProjectPaths, paths


# ============================================================
# Не варьируемые константы ГОСТ 7.32
# ============================================================

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(14)         # основной текст
FONT_SIZE_PAGE_NUMBER = Pt(11)  # номер страницы
FONT_SIZE_TITLE_LARGE = Pt(18)  # вид работы и тема на титульнике
LINE_SPACING_BODY = 1.5
FIRST_LINE_INDENT = Cm(1.25)    # абзацный отступ
PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)


# ============================================================
# Профиль университета — всё, что варьируется между вузами
# ============================================================

@dataclass
class UniversityProfile:
    """Параметры оформления, специфичные для вуза.

    Дефолты соответствуют ГОСТ 7.32-2017 в наиболее распространённом
    толковании. Конкретные вузы (ИТМО, МГУ, СПбГУ и т.д.) переопределяют
    нужные поля. Создавайте свой профиль для своего вуза или принимайте
    PR-ы с новыми константами в этом модуле.
    """
    # Поля титульного листа (мм)
    title_margin_left: float = 30
    title_margin_right: float = 15
    title_margin_top: float = 20
    title_margin_bottom: float = 20
    # Поля основного текста (мм)
    body_margin_left: float = 30
    body_margin_right: float = 15
    body_margin_top: float = 20
    body_margin_bottom: float = 20
    # Кегли заголовков (пт)
    heading_size_h1: int = 16
    heading_size_h2: int = 14
    heading_size_h3: int = 14
    # Выравнивание заголовков: "center" | "left"
    heading_align_h1: str = "center"
    heading_align_h2: str = "left"
    heading_align_h3: str = "left"
    # Поведение h1
    h1_uppercase: bool = True
    h1_new_page: bool = True
    # Заголовок оглавления — ГОСТ 7.32-2017 предписывает «СОДЕРЖАНИЕ»,
    # но многие вузы (ИТМО) используют «ОГЛАВЛЕНИЕ».
    toc_title: str = "СОДЕРЖАНИЕ"
    # Поля титульника, относящиеся к вузу — fallback'и для TitleConfig
    ministry: str = "Министерство науки и высшего образования Российской Федерации"
    university_full: str = ""
    university_short: str = ""
    faculty: str = ""
    city: str = ""


# ---- Встроенные пресеты ----

GOST_PROFILE = UniversityProfile()
"""Чистый ГОСТ 7.32-2017. Поля университета не заданы — должны
прийти из TitleConfig либо из вашего собственного профиля."""

ITMO_PROFILE = UniversityProfile(
    title_margin_left=30,
    title_margin_right=20,
    title_margin_top=10,
    title_margin_bottom=10,
    body_margin_left=30,
    body_margin_right=10,
    body_margin_top=20,
    body_margin_bottom=20,
    heading_size_h1=16,
    heading_size_h2=15,
    heading_size_h3=14,
    heading_align_h1="center",
    heading_align_h2="center",
    heading_align_h3="center",
    toc_title="ОГЛАВЛЕНИЕ",
    university_full=(
        "федеральное государственное автономное "
        "образовательное учреждение высшего образования"
    ),
    university_short="«Национальный исследовательский университет ИТМО»",
    faculty="Факультет программной инженерии и компьютерной техники",
    city="Санкт-Петербург",
)
"""Университет ИТМО, ФПИиКТ. Используется как дефолт, чтобы не ломать
существующие лабораторные."""

DEFAULT_PROFILE = ITMO_PROFILE


# ============================================================
# Ошибка валидации
# ============================================================

class GostValidationError(RuntimeError):
    """Поднимается из Report.save() когда сгенерированный .docx не проходит
    автоматическую проверку ГОСТ.

    Текст исключения — единственная prose-инструкция про валидацию во всём
    скилле. Она появляется в shell-выводе только когда модель уже накосячила.
    Поэтому фразировка имеет значение: явно сказано «перегенерируй», «не
    подавляй», «не используй python-docx напрямую» — это именно те ошибочные
    реакции, которые надо превентить.
    """


# ============================================================
# Конфигурация титульного листа (контент, не оформление)
# ============================================================

@dataclass
class TitleConfig:
    """Содержимое титульного листа.

    Обязательные: work_type, topic, student_name, student_group, year.
    Поля университета (university_*, faculty, city, ministry) можно оставить
    пустыми — тогда они подтянутся из активного UniversityProfile.
    """
    # --- Обязательное ---
    work_type: str
    topic: str
    student_name: str
    student_group: str
    year: str

    # --- Опциональное ---
    work_number: str = ""
    variant: str = ""
    teacher_name: str = ""
    teacher_label: str = "Проверил"        # "Руководитель" для ВКР/курсовых
    teacher_degree: str = ""
    teacher_position: str = ""

    # --- Переопределение полей профиля (пусто = взять из профиля) ---
    city: str = ""
    ministry: str = ""
    university_full: str = ""
    university_short: str = ""
    faculty: str = ""


# ============================================================
# Низкоуровневые помощники работы с XML python-docx
# ============================================================

def _set_run_font(run, *, size=FONT_SIZE_BODY, bold=False, italic=False,
                  underline=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _ensure_normal_style(doc: _Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE_BODY
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING_BODY
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def _add_page_number_to_footer(section, *, hide_on_first=True):
    if hide_on_first:
        section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False

    for p in list(footer.paragraphs):
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run()
    _set_run_font(run, size=FONT_SIZE_PAGE_NUMBER)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)

    if hide_on_first:
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        for p in list(first_footer.paragraphs):
            p.clear()


def _set_section_a4(section):
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT


def _apply_margins(section, left, right, top, bottom):
    section.left_margin = Mm(left)
    section.right_margin = Mm(right)
    section.top_margin = Mm(top)
    section.bottom_margin = Mm(bottom)


def _align_const(name: str):
    return WD_ALIGN_PARAGRAPH.CENTER if name == "center" else WD_ALIGN_PARAGRAPH.LEFT


def _configure_heading_style(doc: _Document, level: int, size_pt: int,
                             align: str):
    style = doc.styles[f"Heading {level}"]
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)

    pf = style.paragraph_format
    pf.alignment = _align_const(align)
    pf.line_spacing = LINE_SPACING_BODY
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True


# ============================================================
# Санитайзер пользовательской прозы
# ============================================================
#
# Любой текст, который агент передаёт в text/task/numbered/bullet/h1-h3 и
# в аргумент caption у figure/table, проходит через _sanitize_prose. Длинное
# и среднее тире — типичный AI-маркер и прямо запрещены writing-style блоком
# в SKILL.md, но модели всё равно их иногда вставляют. Подстраховываемся
# на уровне библиотеки: " — " / " – " между словами становится ", ",
# одиночные тире (например в диапазонах "1—5") — обычным дефисом.
#
# Авто-префиксы "Рисунок N — Описание" и "Таблица N — Описание" формирует
# сама библиотека — там тире обязательно по ГОСТ и не трогается.

_PROSE_DASH_BETWEEN_WORDS = re.compile(r"\s+[—–]\s+")
_PROSE_DASH_REMAINING = re.compile(r"[—–]")


def _sanitize_prose(text):
    if not text:
        return text
    text = _PROSE_DASH_BETWEEN_WORDS.sub(", ", text)
    text = _PROSE_DASH_REMAINING.sub("-", text)
    return text


# ============================================================
# LaTeX → OMML
# ============================================================
#
# Word хранит формулы в OMML (Office Math, namespace
# http://schemas.openxmlformats.org/officeDocument/2006/math, prefix m:),
# а не в MathML. Поэтому путь такой:
#
#     LaTeX  --[latex2mathml]-->  MathML  --[наш walker]-->  OMML
#
# latex2mathml делает самую противную часть (нерегулярная грамматика LaTeX),
# а MathML→OMML это прямой обход дерева: теги мапятся почти 1:1.
#
# Сама строка LaTeX **не** проходит через _sanitize_prose — иначе
# `\text{1—5}` или `[a—b]` сломались бы. Санируется только пользовательская
# проза в where=... (это обычный русский текст, к нему применимы общие
# правила: никаких длинных тире).

_MATHML_NS = "http://www.w3.org/1998/Math/MathML"
_M_PREFIX = "{" + _MATHML_NS + "}"

# Реляционные операторы — границы тела N-ary (после `=`, `<`, `≥` и т.п.
# идёт уже не подынтегральное выражение, а вторая часть равенства). При
# сборе body для <m:e> упираемся в эти символы и останавливаемся.
_NARY_BODY_TERMINATORS = {
    "=", "<", ">", "≤", "≥", "≠", "≈", "≡",
    "≪", "≫", "⇔", "⇒", "↔", "→", "←", "∝",
}


# N-ary операторы: символ → размещение пределов. undOvr = пределы
# сверху/снизу (∑, ∏), subSup = справа от знака (∫, ∮).
_NARY_OPS = {
    "∑": "undOvr",  # ∑
    "∏": "undOvr",  # ∏
    "∐": "undOvr",  # ∐
    "⋃": "undOvr",  # ⋃
    "⋂": "undOvr",  # ⋂
    "⨂": "undOvr",  # ⨂
    "⨁": "undOvr",  # ⨁
    "⨀": "undOvr",  # ⨀
    "∫": "subSup",  # ∫
    "∬": "subSup",  # ∬
    "∭": "subSup",  # ∭
    "∮": "subSup",  # ∮
}

# Акценты: \bar, \hat, \vec, \tilde, \dot, \ddot, \check, \acute, \grave, \breve.
# latex2mathml выдаёт spacing-формы (¯ ^ ~ ¨ ´ ` ˘ ˙ ˇ →), но Word в <m:acc>
# умеет красиво рисовать только combining-формы (U+0300-U+030C, U+20D7).
# Spacing-знак рисуется у baseline и пересекает букву; combining-знак — сверху.
# При эмите OMML заменяем spacing → combining через _ACCENT_NORMALIZE.

_ACCENT_NORMALIZE = {
    "¯": "̄",  # MACRON → COMBINING MACRON (\bar)
    "^": "̂",  # CIRCUMFLEX → COMBINING CIRCUMFLEX (\hat)
    "~": "̃",  # TILDE → COMBINING TILDE (\tilde)
    "¨": "̈",  # DIAERESIS → COMBINING DIAERESIS (\ddot)
    "´": "́",  # ACUTE → COMBINING ACUTE (\acute)
    "`": "̀",  # GRAVE → COMBINING GRAVE (\grave)
    "˘": "̆",  # BREVE → COMBINING BREVE (\breve)
    "˙": "̇",  # DOT ABOVE → COMBINING DOT ABOVE (\dot)
    "ˇ": "̌",  # CARON → COMBINING CARON (\check)
    "→": "⃗",  # RIGHTWARDS ARROW → COMBINING RIGHT ARROW ABOVE (\vec)
}

_ACCENT_CHARS = set(_ACCENT_NORMALIZE.keys()) | set(_ACCENT_NORMALIZE.values())


def _ml_local(tag: str) -> str:
    """Имя MathML-тега без namespace-префикса."""
    if tag.startswith(_M_PREFIX):
        return tag[len(_M_PREFIX):]
    return tag


def _omml(tag: str) -> OxmlElement:
    """OMML-элемент с префиксом m: (math namespace зарегистрирован в python-docx)."""
    return OxmlElement(f"m:{tag}")


def _set_mval(el: OxmlElement, val: str) -> None:
    el.set(qn("m:val"), val)


def _omml_run(text: str, *, plain: bool = False) -> OxmlElement:
    """<m:r>[<m:rPr><m:sty m:val="p"/></m:rPr>]<m:t>text</m:t></m:r>.

    plain=True ставит «прямой» стиль (для чисел, операторов, многобуквенных
    идентификаторов вроде sin/log). Для одиночных букв — курсив (дефолт OMML).
    """
    r = _omml("r")
    if plain:
        rPr = _omml("rPr")
        sty = _omml("sty")
        _set_mval(sty, "p")
        rPr.append(sty)
        r.append(rPr)
    t = _omml("t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _omml_wrap(tag: str, children: List[OxmlElement]) -> OxmlElement:
    """<m:tag>...children...</m:tag>"""
    wrapper = _omml(tag)
    for c in children:
        wrapper.append(c)
    return wrapper


def _is_nary_op(node) -> Optional[str]:
    """Вернёт N-ary символ если узел — это <mo>∑</mo> / <mo>∫</mo> и т.п."""
    if _ml_local(node.tag) != "mo":
        return None
    text = (node.text or "").strip()
    return text if text in _NARY_OPS else None


def _build_nary(chr_text: str,
                sub_children: List[OxmlElement],
                sup_children: List[OxmlElement],
                *,
                e_children: Optional[List[OxmlElement]] = None,
                hide_sub: bool = False,
                hide_sup: bool = False) -> OxmlElement:
    """<m:nary> для N-ary оператора с пределами и подынтегральным <m:e>.

    e_children — содержимое тела (что идёт после знака суммы/интеграла).
    Если None или пустой список, Word нарисует placeholder-квадрат на месте
    тела — поэтому собирать body нужно на уровне mrow с помощью lookahead.
    """
    nary = _omml("nary")
    naryPr = _omml("naryPr")
    chr_el = _omml("chr")
    _set_mval(chr_el, chr_text)
    naryPr.append(chr_el)
    limLoc = _omml("limLoc")
    _set_mval(limLoc, _NARY_OPS.get(chr_text, "subSup"))
    naryPr.append(limLoc)
    if hide_sub:
        sh = _omml("subHide")
        _set_mval(sh, "1")
        naryPr.append(sh)
    if hide_sup:
        sh = _omml("supHide")
        _set_mval(sh, "1")
        naryPr.append(sh)
    nary.append(naryPr)
    nary.append(_omml_wrap("sub", sub_children))
    nary.append(_omml_wrap("sup", sup_children))
    nary.append(_omml_wrap("e", e_children or []))
    return nary


def _extract_nary_info(node):
    """Если node это msub/msup/msubsup/munder/mover/munderover, базой которой
    является N-ary оператор (∑, ∫, ∏, …) — вернёт кортеж
    (chr_text, sub_children, sup_children, hide_sub, hide_sup).
    Иначе None.
    """
    tag = _ml_local(node.tag)
    kids = list(node)
    if tag in ("msup", "mover") and len(kids) >= 2:
        chr_text = _is_nary_op(kids[0])
        if chr_text:
            return (chr_text, [], _walk_mathml(kids[1]), True, False)
    if tag in ("msub", "munder") and len(kids) >= 2:
        chr_text = _is_nary_op(kids[0])
        if chr_text:
            return (chr_text, _walk_mathml(kids[1]), [], False, True)
    if tag in ("msubsup", "munderover") and len(kids) >= 3:
        chr_text = _is_nary_op(kids[0])
        if chr_text:
            return (chr_text,
                    _walk_mathml(kids[1]),
                    _walk_mathml(kids[2]),
                    False, False)
    return None


def _is_body_terminator(node) -> bool:
    """Останавливает сбор body N-ary оператора. Реляционные операторы (=, ≤,
    ⇒ и т.п.) разделяют интегранд от правой части уравнения."""
    if _ml_local(node.tag) != "mo":
        return False
    text = (node.text or "").strip()
    return text in _NARY_BODY_TERMINATORS


def _walk_with_nary(children, start: int, *, stop_at_terminator: bool):
    """Обход списка MathML-детей с lookahead для N-ary операторов.

    Возвращает (omml_elements, next_index) — индекс на первый элемент,
    который не был поглощён (терминатор либо конец списка).

    stop_at_terminator=True: останавливаемся при первом теле-терминаторе
        (=, ≤, ≠, …) и НЕ съедаем его — он остаётся для caller'а. Это
        режим сбора body внутри N-ary: правая часть равенства не
        принадлежит подынтегральному выражению.
    stop_at_terminator=False: терминаторы рендерятся как обычные mo,
        потребляются. Режим для верхнего mrow.

    Рекурсивный — вложенные N-ary (∑∑..., ∫∑..., etc.) обрабатываются
    корректно: внутренний оператор поглощает свой body раньше, чем
    внешний продолжает сбор.
    """
    result: List[OxmlElement] = []
    j = start
    while j < len(children):
        child = children[j]
        if stop_at_terminator and _is_body_terminator(child):
            break
        nary_info = _extract_nary_info(child)
        if nary_info:
            chr_text, sub_kids, sup_kids, hide_sub, hide_sup = nary_info
            inner_body, next_j = _walk_with_nary(
                children, j + 1, stop_at_terminator=True
            )
            result.append(_build_nary(
                chr_text, sub_kids, sup_kids,
                e_children=inner_body,
                hide_sub=hide_sub, hide_sup=hide_sup,
            ))
            j = next_j
            continue
        result.extend(_walk_mathml(child))
        j += 1
    return result, j


def _walk_mathml(node) -> List[OxmlElement]:
    """Рекурсивный обход MathML-узла, возвращает список OMML-элементов.

    Возвращаем именно список (а не один узел), потому что mrow/mstyle
    плющатся в плоский список детей при подстановке в обёртки типа
    <m:e>, <m:num>, <m:sup>.
    """
    tag = _ml_local(node.tag)

    # Контейнеры — плющим в плоский список с lookahead для N-ary тел.
    if tag in ("math", "mstyle", "mrow", "semantics", "annotation"):
        children = [c for c in node if _ml_local(c.tag) != "annotation"]
        result, _ = _walk_with_nary(children, 0, stop_at_terminator=False)
        return result

    if tag == "mi":
        text = (node.text or "").strip()
        if not text:
            return []
        # Многобуквенные идентификаторы (sin, log, lim, exp) — прямые;
        # одиночные буквы (включая греческие) — курсив (дефолт OMML).
        # mathvariant="normal" принудительно делает прямой шрифт.
        plain = len(text) > 1 or node.get("mathvariant") == "normal"
        return [_omml_run(text, plain=plain)]

    if tag in ("mn", "mo", "mtext"):
        text = node.text or ""
        if not text or not text.strip():
            # Пустые mo (вокруг скобок и т.п.) — часто нерелевантны
            if tag == "mtext":
                return [_omml_run(text, plain=True)] if text else []
            return []
        return [_omml_run(text, plain=True)]

    if tag == "mspace":
        return [_omml_run(" ", plain=True)]

    if tag == "mfrac":
        kids = list(node)
        if len(kids) < 2:
            return []
        f = _omml("f")
        f.append(_omml_wrap("num", _walk_mathml(kids[0])))
        f.append(_omml_wrap("den", _walk_mathml(kids[1])))
        return [f]

    if tag == "msup":
        kids = list(node)
        if len(kids) < 2:
            return []
        # latex2mathml в inline-режиме оборачивает \sum^{n}, \int^{n} и т.п.
        # в msup/msub/msubsup — детектируем N-ary базу до обычного sSup.
        nary_chr = _is_nary_op(kids[0])
        if nary_chr:
            return [_build_nary(nary_chr, [], _walk_mathml(kids[1]),
                                hide_sub=True)]
        s = _omml("sSup")
        s.append(_omml_wrap("e", _walk_mathml(kids[0])))
        s.append(_omml_wrap("sup", _walk_mathml(kids[1])))
        return [s]

    if tag == "msub":
        kids = list(node)
        if len(kids) < 2:
            return []
        nary_chr = _is_nary_op(kids[0])
        if nary_chr:
            return [_build_nary(nary_chr, _walk_mathml(kids[1]), [],
                                hide_sup=True)]
        s = _omml("sSub")
        s.append(_omml_wrap("e", _walk_mathml(kids[0])))
        s.append(_omml_wrap("sub", _walk_mathml(kids[1])))
        return [s]

    if tag == "msubsup":
        kids = list(node)
        if len(kids) < 3:
            return []
        nary_chr = _is_nary_op(kids[0])
        if nary_chr:
            return [_build_nary(nary_chr,
                                _walk_mathml(kids[1]),
                                _walk_mathml(kids[2]))]
        s = _omml("sSubSup")
        s.append(_omml_wrap("e", _walk_mathml(kids[0])))
        s.append(_omml_wrap("sub", _walk_mathml(kids[1])))
        s.append(_omml_wrap("sup", _walk_mathml(kids[2])))
        return [s]

    if tag == "msqrt":
        rad = _omml("rad")
        radPr = _omml("radPr")
        degHide = _omml("degHide")
        _set_mval(degHide, "1")
        radPr.append(degHide)
        rad.append(radPr)
        rad.append(_omml("deg"))
        e_kids: List[OxmlElement] = []
        for child in node:
            e_kids.extend(_walk_mathml(child))
        rad.append(_omml_wrap("e", e_kids))
        return [rad]

    if tag == "mroot":
        kids = list(node)
        if len(kids) < 2:
            return []
        rad = _omml("rad")
        rad.append(_omml_wrap("deg", _walk_mathml(kids[1])))
        rad.append(_omml_wrap("e", _walk_mathml(kids[0])))
        return [rad]

    if tag == "mover":
        # Семантика MathML: kids[0] = база, kids[1] = надстрочный (overscript).
        # Это либо акцент (\bar, \hat, \vec) либо предел над оператором.
        kids = list(node)
        if len(kids) < 2:
            return []
        base, over = kids[0], kids[1]

        # Акцент: либо явный accent="true", либо overscript это <mo> с
        # одним маркером из _ACCENT_CHARS.
        over_text = (over.text or "").strip() if _ml_local(over.tag) == "mo" else ""
        is_accent = (
            node.get("accent") == "true"
            or (len(over_text) <= 2 and over_text in _ACCENT_CHARS)
        )
        if is_accent and over_text:
            # Spacing-знак (¯ ^ ~ ¨ ´ ` ˘ ˙ ˇ →) перевести в combining-форму,
            # иначе Word рисует его на baseline и накладывает поверх буквы.
            accent_chr = _ACCENT_NORMALIZE.get(over_text, over_text)
            acc = _omml("acc")
            accPr = _omml("accPr")
            chr_el = _omml("chr")
            _set_mval(chr_el, accent_chr)
            accPr.append(chr_el)
            acc.append(accPr)
            acc.append(_omml_wrap("e", _walk_mathml(base)))
            return [acc]

        # N-ary оператор с одним только верхним пределом (редко, но возможно).
        nary_chr = _is_nary_op(base)
        if nary_chr:
            return [_build_nary(nary_chr, [], _walk_mathml(over),
                                hide_sub=True)]

        # Иначе: оператор с верхним пределом → m:limUpp.
        lu = _omml("limUpp")
        lu.append(_omml_wrap("e", _walk_mathml(base)))
        lu.append(_omml_wrap("lim", _walk_mathml(over)))
        return [lu]

    if tag == "munder":
        # kids[0] = база, kids[1] = подстрочный (underscript).
        kids = list(node)
        if len(kids) < 2:
            return []
        base, under = kids[0], kids[1]

        nary_chr = _is_nary_op(base)
        if nary_chr:
            return [_build_nary(nary_chr, _walk_mathml(under), [],
                                hide_sup=True)]

        ll = _omml("limLow")
        ll.append(_omml_wrap("e", _walk_mathml(base)))
        ll.append(_omml_wrap("lim", _walk_mathml(under)))
        return [ll]

    if tag == "munderover":
        # kids[0] = база, kids[1] = under, kids[2] = over.
        kids = list(node)
        if len(kids) < 3:
            return []
        base, under, over = kids[0], kids[1], kids[2]

        nary_chr = _is_nary_op(base)
        if nary_chr:
            return [_build_nary(nary_chr,
                                _walk_mathml(under),
                                _walk_mathml(over))]

        # Generic fallback: вложенные limUpp(limLow(...)).
        ll = _omml("limLow")
        ll.append(_omml_wrap("e", _walk_mathml(base)))
        ll.append(_omml_wrap("lim", _walk_mathml(under)))
        lu = _omml("limUpp")
        lu.append(_omml_wrap("e", [ll]))
        lu.append(_omml_wrap("lim", _walk_mathml(over)))
        return [lu]

    if tag == "mfenced":
        open_chr = node.get("open", "(")
        close_chr = node.get("close", ")")
        d = _omml("d")
        dPr = _omml("dPr")
        if open_chr != "(":
            beg = _omml("begChr")
            _set_mval(beg, open_chr)
            dPr.append(beg)
        if close_chr != ")":
            end = _omml("endChr")
            _set_mval(end, close_chr)
            dPr.append(end)
        if list(dPr):
            d.append(dPr)
        e_kids = []
        for child in node:
            e_kids.extend(_walk_mathml(child))
        d.append(_omml_wrap("e", e_kids))
        return [d]

    if tag == "mtable":
        m_el = _omml("m")
        for row_node in node:
            if _ml_local(row_node.tag) != "mtr":
                continue
            mr = _omml("mr")
            for cell_node in row_node:
                if _ml_local(cell_node.tag) != "mtd":
                    continue
                cell_kids: List[OxmlElement] = []
                for child in cell_node:
                    cell_kids.extend(_walk_mathml(child))
                mr.append(_omml_wrap("e", cell_kids))
            m_el.append(mr)
        return [m_el]

    # Неизвестный тег — рекурсивно обрабатываем детей, не падаем
    result = []
    for child in node:
        result.extend(_walk_mathml(child))
    return result


def _latex_to_omath(latex: str) -> OxmlElement:
    """LaTeX-строка → <m:oMath> готовый к вставке в параграф."""
    try:
        import latex2mathml.converter as _l2m
    except ImportError as e:
        raise ImportError(
            "r.formula() requires the 'latex2mathml' package. "
            "If you launched the script through scripts/ensure_env.py, the "
            "venv setup must have failed. Otherwise install manually: "
            "pip install latex2mathml"
        ) from e

    mml_str = _l2m.convert(latex)
    tree = _ET.fromstring(mml_str)
    children = _walk_mathml(tree)
    omath = _omml("oMath")
    for c in children:
        omath.append(c)
    return omath


# ============================================================
# Класс отчёта — высокоуровневый API
# ============================================================

class Report:
    """Главный класс. Собирает документ из вызовов h1/text/code/figure/...

    Создание:
        r = Report(TitleConfig(...))                    # ITMO_PROFILE по умолчанию
        r = Report(TitleConfig(...), profile=GOST_PROFILE)
        r = Report(TitleConfig(...), profile=UniversityProfile(...))

    Контент: см. методы toc/h1/h2/h3/text/task/code/figure/table/numbered/bullet.
    Сохранение: r.save("report.docx").
    """

    def __init__(self, title: TitleConfig,
                 profile: UniversityProfile = DEFAULT_PROFILE,
                 *,
                 project_root: Optional[Union[str, Path]] = None):
        self._doc = Document()
        self._title = title
        self._profile = profile
        self._figure_counter = 0
        self._table_counter = 0
        self._formula_counter = 0
        self._just_broke_page = False
        self._next_num_id = 100

        # Авто-резолв путей: если project_root не задан — paths() обходит
        # стек вверх до user-скрипта, оттуда ищет маркер проекта (.git,
        # Makefile, pyproject.toml, .claude). Используется figure() для
        # резолва относительных имён файлов и save() для дефолтного пути.
        if project_root is None:
            self._paths = paths()
        else:
            self._paths = paths(Path(project_root))

        # 1. Базовые стили
        _ensure_normal_style(self._doc)
        _configure_heading_style(self._doc, 1,
                                 profile.heading_size_h1,
                                 profile.heading_align_h1)
        _configure_heading_style(self._doc, 2,
                                 profile.heading_size_h2,
                                 profile.heading_align_h2)
        _configure_heading_style(self._doc, 3,
                                 profile.heading_size_h3,
                                 profile.heading_align_h3)

        # 2. Секция титульника (первая)
        section = self._doc.sections[0]
        _set_section_a4(section)
        _apply_margins(section,
                       profile.title_margin_left, profile.title_margin_right,
                       profile.title_margin_top, profile.title_margin_bottom)
        _add_page_number_to_footer(section, hide_on_first=True)
        self._build_title_page()

        # 3. Новая секция для основного текста
        body_section = self._doc.add_section(WD_SECTION.NEW_PAGE)
        _set_section_a4(body_section)
        _apply_margins(body_section,
                       profile.body_margin_left, profile.body_margin_right,
                       profile.body_margin_top, profile.body_margin_bottom)
        _add_page_number_to_footer(body_section, hide_on_first=False)
        self._just_broke_page = True

    # --------------------------------------------------------
    # Внутреннее: построение титульника
    # --------------------------------------------------------

    def _resolve(self, attr: str) -> str:
        """TitleConfig.<attr> если задано, иначе UniversityProfile.<attr>."""
        value = getattr(self._title, attr, "") or ""
        if value:
            return value
        return getattr(self._profile, attr, "") or ""

    def _add_paragraph(self, text="", *, align=WD_ALIGN_PARAGRAPH.CENTER,
                       size=FONT_SIZE_BODY, bold=False, italic=False,
                       underline=False, left_indent=None,
                       first_line_indent=None):
        p = self._doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if left_indent is not None:
            pf.left_indent = left_indent
        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent
        if text:
            run = p.add_run(text)
            _set_run_font(run, size=size, bold=bold, italic=italic,
                          underline=underline)
        return p

    def _add_runs_paragraph(self, runs, *, align=WD_ALIGN_PARAGRAPH.CENTER,
                            left_indent=None):
        p = self._doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if left_indent is not None:
            pf.left_indent = left_indent
        for r in runs:
            run = p.add_run(r["text"])
            _set_run_font(
                run,
                size=r.get("size", FONT_SIZE_BODY),
                bold=r.get("bold", False),
                italic=r.get("italic", False),
                underline=r.get("underline", False),
            )
        return p

    def _add_bottom_anchored_block(self, lines):
        """Borderless single-column таблица, плавающе прижатая к нижнему
        полю текущей страницы. Каждый элемент ``lines`` — отдельная строка
        (центр, TNR 14). Используется для города/года на титульнике, чтобы
        они оставались на 1-й странице независимо от длины шапки/темы/ФИО.

        Реализация — через OOXML ``w:tblpPr`` (плавающее позиционирование):
        ``vertAnchor=margin`` + ``tblpYSpec=bottom`` прижимает низ таблицы к
        нижнему полю; ``horzAnchor=margin`` + ``tblpXSpec=center`` центрирует
        по горизонтали внутри полей. python-docx высокоуровневого API для
        этого не даёт, поэтому правим XML напрямую.
        """
        table = self._doc.add_table(rows=len(lines), cols=1)
        table.autofit = True

        tbl = table._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblpPr = OxmlElement("w:tblpPr")
        tblpPr.set(qn("w:vertAnchor"), "margin")
        tblpPr.set(qn("w:horzAnchor"), "margin")
        tblpPr.set(qn("w:tblpYSpec"), "bottom")
        tblpPr.set(qn("w:tblpXSpec"), "center")
        tblpPr.set(qn("w:leftFromText"), "0")
        tblpPr.set(qn("w:rightFromText"), "0")
        tblpPr.set(qn("w:topFromText"), "0")
        tblpPr.set(qn("w:bottomFromText"), "0")
        tblPr.append(tblpPr)

        tblOverlap = OxmlElement("w:tblOverlap")
        tblOverlap.set(qn("w:val"), "never")
        tblPr.append(tblOverlap)

        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        for row, text in zip(table.rows, lines):
            cell = row.cells[0]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING_BODY
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            run = p.add_run(text)
            _set_run_font(run, size=FONT_SIZE_BODY)

        return table

    def _build_title_page(self):
        cfg = self._title

        # Шапка: министерство → университет → факультет
        ministry = self._resolve("ministry")
        if ministry:
            self._add_paragraph(ministry)
        university_full = self._resolve("university_full")
        if university_full:
            self._add_paragraph(university_full)
        university_short = self._resolve("university_short")
        if university_short:
            self._add_paragraph(university_short)
        faculty = self._resolve("faculty")
        if faculty:
            self._add_paragraph()
            self._add_paragraph(faculty, italic=True)

        # Отступ к центру листа
        for _ in range(6):
            self._add_paragraph()

        # Вид работы (крупно, полужирно)
        work_type_text = cfg.work_type
        if cfg.work_number:
            work_type_text = f"{work_type_text} {cfg.work_number}".strip()
        self._add_paragraph(work_type_text, size=FONT_SIZE_TITLE_LARGE,
                            bold=True)

        # Тема (крупно, без кавычек)
        if cfg.topic:
            self._add_paragraph(cfg.topic, size=FONT_SIZE_TITLE_LARGE)

        # Вариант
        if cfg.variant:
            self._add_paragraph()
            self._add_paragraph(f"Вариант №{cfg.variant}",
                                size=FONT_SIZE_TITLE_LARGE)

        # Отступ к блоку студент/преподаватель
        for _ in range(4):
            self._add_paragraph()

        right_block_indent = Cm(9)

        self._add_runs_paragraph(
            [{"text": f"Группа: {cfg.student_group}", "underline": True}],
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=right_block_indent,
        )
        self._add_runs_paragraph(
            [
                {"text": "Выполнил", "underline": True},
                {"text": f": {cfg.student_name}"},
            ],
            align=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=right_block_indent,
        )

        if cfg.teacher_name:
            self._add_paragraph(left_indent=right_block_indent,
                                align=WD_ALIGN_PARAGRAPH.LEFT)
            self._add_runs_paragraph(
                [
                    {"text": cfg.teacher_label, "underline": True},
                    {"text": ":"},
                ],
                align=WD_ALIGN_PARAGRAPH.LEFT,
                left_indent=right_block_indent,
            )
            parts = []
            if cfg.teacher_degree:
                parts.append(cfg.teacher_degree)
            if cfg.teacher_position:
                parts.append(cfg.teacher_position)
            parts.append(cfg.teacher_name)
            self._add_paragraph(
                " ".join(parts),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                left_indent=right_block_indent,
            )

        # Город и год — прижаты к нижнему полю плавающей borderless-таблицей,
        # чтобы не уезжать на page 2 при длинной шапке/теме/ФИО.
        city = self._resolve("city")
        footer_lines = []
        if city:
            footer_lines.append(city)
        footer_lines.append(cfg.year)
        self._add_bottom_anchored_block(footer_lines)

    # --------------------------------------------------------
    # Публичный API: контент основного текста
    # --------------------------------------------------------

    def toc(self):
        """Вставляет автоматическое оглавление (поле Word TOC). Заголовок
        — `profile.toc_title` (по умолчанию «СОДЕРЖАНИЕ» в GOST_PROFILE,
        «ОГЛАВЛЕНИЕ» в ITMO_PROFILE).

        После открытия файла в Word: ПКМ по полю → «Обновить поле».
        """
        heading = self._doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.line_spacing = LINE_SPACING_BODY
        heading.paragraph_format.space_after = Pt(12)
        run = heading.add_run(self._profile.toc_title)
        _set_run_font(run, size=Pt(self._profile.heading_size_h1), bold=True)

        p = self._doc.add_paragraph()
        run = p.add_run()
        _set_run_font(run)

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        placeholder_run = OxmlElement("w:r")
        placeholder_text = OxmlElement("w:t")
        placeholder_text.text = ("Оглавление будет сформировано автоматически "
                                 "при обновлении поля (ПКМ → «Обновить поле»)")
        placeholder_run.append(placeholder_text)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._element.append(fld_begin)
        run._element.append(instr)
        run._element.append(fld_separate)
        run._element.append(placeholder_run)
        run._element.append(fld_end)

        self.page_break()

    def h1(self, text: str):
        """Заголовок 1 уровня. Поведение управляется профилем:
        h1_new_page (новая страница), h1_uppercase (ВЕРХНИЙ регистр).
        """
        if self._profile.h1_new_page and not self._just_broke_page:
            self.page_break()
        p = self._doc.add_paragraph(style="Heading 1")
        text = _sanitize_prose(text)
        text_to_render = text.upper() if self._profile.h1_uppercase else text
        run = p.add_run(text_to_render)
        _set_run_font(run, size=Pt(self._profile.heading_size_h1), bold=True)
        self._just_broke_page = False

    def h2(self, text: str):
        p = self._doc.add_paragraph(style="Heading 2")
        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run, size=Pt(self._profile.heading_size_h2), bold=True)

    def h3(self, text: str):
        p = self._doc.add_paragraph(style="Heading 3")
        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run, size=Pt(self._profile.heading_size_h3), bold=True)

    def text(self, text: str, *, bold=False, italic=False):
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.first_line_indent = FIRST_LINE_INDENT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run, bold=bold, italic=italic)

    def task(self, text: str):
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.first_line_indent = FIRST_LINE_INDENT
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run, bold=True)

    def code(self, code: str):
        for line in code.split("\n"):
            p = self._doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.line_spacing = 1.0
            pf.first_line_indent = Pt(0)
            pf.left_indent = Cm(0.5)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            run = p.add_run(line if line else " ")
            run.font.name = "Courier New"
            run.font.size = Pt(11)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rfonts.set(qn(attr), "Courier New")

    def _resolve_figure_path(self, image_path: Union[str, Path]) -> Path:
        """Абсолютный путь — без изменений; относительный — от
        `self._paths.figures`. Бросает FileNotFoundError с обоими путями
        (input + resolved) если файл не существует.

        BC note: ранее относительные пути резолвились от os.getcwd(); теперь
        от <project>/docs/figures/. Реальные user-скрипты передавали
        абсолютные пути через Path(__file__).parent, поэтому wild-impact
        близок к нулю. Подробнее — CHANGELOG v0.7.0.
        """
        p = Path(image_path)
        if p.is_absolute():
            resolved = p
        else:
            resolved = self._paths.figures / p
        if not resolved.exists():
            raise FileNotFoundError(
                f"figure not found: input={image_path!r}, resolved={resolved}"
            )
        return resolved

    def figure(self, image_path: Union[str, Path], caption: str, *,
               width_cm: Optional[float] = None):
        """Вставляет рисунок с автоматической подписью.

        Ширина картинки **всегда** ограничивается печатной областью страницы
        (A4 минус левое и правое поля активного профиля — обычно ~17 см для
        ITMO и ~16.5 см для GOST). Если width_cm не задан, используется
        натуральный размер картинки (с клампом к печатной области, если она
        больше). Если width_cm задан — он уважается, но также клампится.
        Это предотвращает вылет крупных скриншотов за поля.

        Резолв пути:
        - абсолютный путь — без изменений.
        - относительный — резолвится от `<project>/docs/figures/`.
          `r.figure("load.png", ...)` → `<project>/docs/figures/load.png`.
          `r.figure("subdir/load.png", ...)` — тоже от `figures/`.
        Если файл не найден — FileNotFoundError с обоими путями (input + resolved).
        """
        image_path = self._resolve_figure_path(image_path)
        self._figure_counter += 1

        # Печатная область: 210 мм (A4) − левое поле − правое поле, в см
        max_width_cm = (210
                        - self._profile.body_margin_left
                        - self._profile.body_margin_right) / 10.0

        image_path_str = str(image_path)
        if width_cm is None:
            # Прочитать натуральные размеры через python-docx (без PIL-зависимости)
            from docx.image.image import Image as _DocxImage
            img_meta = _DocxImage.from_file(image_path_str)
            natural_width_cm = img_meta.px_width / img_meta.horz_dpi * 2.54
            if natural_width_cm > max_width_cm:
                width_cm = max_width_cm
            # else: натуральный размер влезает, оставляем width_cm=None
        elif width_cm > max_width_cm:
            width_cm = max_width_cm

        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(0)
        run = p.add_run()
        if width_cm is not None:
            run.add_picture(image_path_str, width=Cm(width_cm))
        else:
            run.add_picture(image_path_str)

        cap = self._doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cpf = cap.paragraph_format
        cpf.line_spacing = LINE_SPACING_BODY
        cpf.first_line_indent = Pt(0)
        cpf.space_before = Pt(0)
        cpf.space_after = Pt(6)
        cap_run = cap.add_run(
            f"Рисунок {self._figure_counter} — {_sanitize_prose(caption)}")
        _set_run_font(cap_run)

    def formula(self, latex: str, *, where: Optional[str] = None) -> int:
        """Вставляет формулу из LaTeX как нативное Word-уравнение (OMML).

        Возвращает номер формулы — пригодится для ссылок:
            f1 = r.formula(r"E = mc^2")
            r.text(f"По формуле ({f1}) видно ...")

        Параметр `where` — пояснения переменных, идут отдельным абзацем
        под формулой:
            r.formula(
                r"\\frac{a + b}{c}",
                where="a, b — слагаемые, c — делитель",
            )

        Сама строка `latex` не санируется (LaTeX-синтаксис не должен
        искажаться); `where` проходит через _sanitize_prose как обычная
        проза. Зависит от пакета `latex2mathml` — он ставится автоматически
        через scripts/ensure_env.py.

        Форматирование по ГОСТ 7.32: формула центрирована, номер «(N)»
        прижат к правому краю печатной области на той же строке.
        """
        self._formula_counter += 1
        n = self._formula_counter

        omath = _latex_to_omath(latex)

        # Печатная ширина (A4 минус поля активного профиля), та же формула,
        # что в figure() — даём числу прижаться к правому краю.
        printable_cm = (210
                        - self._profile.body_margin_left
                        - self._profile.body_margin_right) / 10.0

        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.first_line_indent = Pt(0)
        pf.left_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(0) if where else Pt(6)

        # Центр-таб по середине печатной области, правый таб — по правому краю.
        # Layout: TAB(center)<formula>TAB(right)(N)
        pf.tab_stops.add_tab_stop(Cm(printable_cm / 2), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(printable_cm), WD_TAB_ALIGNMENT.RIGHT)

        # Tab → центр
        r_tab1 = OxmlElement("w:r")
        r_tab1.append(OxmlElement("w:tab"))
        p._p.append(r_tab1)

        # Сама формула
        p._p.append(omath)

        # Tab → правый край → "(N)"
        num_run = p.add_run()
        num_run._element.append(OxmlElement("w:tab"))
        num_t = OxmlElement("w:t")
        num_t.text = f"({n})"
        num_run._element.append(num_t)
        _set_run_font(num_run)

        if where:
            wp = self._doc.add_paragraph()
            wp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            wpf = wp.paragraph_format
            wpf.line_spacing = LINE_SPACING_BODY
            wpf.first_line_indent = Pt(0)
            wpf.left_indent = Cm(0)
            wpf.space_before = Pt(0)
            wpf.space_after = Pt(6)
            run = wp.add_run(f"где {_sanitize_prose(where)}")
            _set_run_font(run)

        return n

    def table(self, rows: Sequence[Sequence[str]], caption: str = "",
              *, has_header: bool = True):
        if not rows:
            return
        self._table_counter += 1

        if caption:
            cap = self._doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cpf = cap.paragraph_format
            cpf.line_spacing = LINE_SPACING_BODY
            cpf.first_line_indent = Pt(0)
            cpf.space_before = Pt(6)
            cpf.space_after = Pt(0)
            cap_run = cap.add_run(
                f"Таблица {self._table_counter} — {_sanitize_prose(caption)}")
            _set_run_font(cap_run)

        n_cols = max(len(r) for r in rows)
        table = self._doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Table Grid"
        for i, row_data in enumerate(rows):
            for j in range(n_cols):
                cell = table.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.first_line_indent = Pt(0)
                value = row_data[j] if j < len(row_data) else ""
                run = p.add_run(_sanitize_prose(value))
                _set_run_font(run, bold=(has_header and i == 0))

    def _create_independent_num(self, abstract_num_id: str) -> int:
        numbering = self._doc.part.numbering_part.element
        num_id = self._next_num_id
        self._next_num_id += 1

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), abstract_num_id)
        num.append(abs_ref)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)

        numbering.append(num)
        return num_id

    def _find_abstract_num_for_style(self, style_name: str) -> Optional[str]:
        try:
            style = self._doc.styles[style_name]
        except KeyError:
            return None
        style_pPr = style.element.find(qn("w:pPr"))
        if style_pPr is None:
            return None
        numPr = style_pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        numId_el = numPr.find(qn("w:numId"))
        if numId_el is None:
            return None
        style_num_id = numId_el.get(qn("w:val"))

        numbering = self._doc.part.numbering_part.element
        for num in numbering.findall(qn("w:num")):
            if num.get(qn("w:numId")) == style_num_id:
                abs_ref = num.find(qn("w:abstractNumId"))
                if abs_ref is not None:
                    return abs_ref.get(qn("w:val"))
        return None

    def _add_list_paragraph(self, text: str, num_id: int):
        p = self._doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.space_after = Pt(0)

        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(num_id))
        numPr.append(ilvl)
        numPr.append(numId_el)
        pPr.append(numPr)

        run = p.add_run(_sanitize_prose(text))
        _set_run_font(run)

    def numbered(self, items):
        if isinstance(items, str):
            items = [items]
        if not items:
            return
        try:
            _ = self._doc.part.numbering_part
        except (AttributeError, KeyError):
            tmp = self._doc.add_paragraph(style="List Number")
            tmp._element.getparent().remove(tmp._element)
        abstract_id = self._find_abstract_num_for_style("List Number")
        if abstract_id is None:
            for item in items:
                p = self._doc.add_paragraph(style="List Number")
                pf = p.paragraph_format
                pf.line_spacing = LINE_SPACING_BODY
                pf.space_after = Pt(0)
                run = p.add_run(_sanitize_prose(item))
                _set_run_font(run)
            return
        num_id = self._create_independent_num(abstract_id)
        for item in items:
            self._add_list_paragraph(item, num_id)

    def bullet(self, items):
        if isinstance(items, str):
            items = [items]
        if not items:
            return
        try:
            _ = self._doc.part.numbering_part
        except (AttributeError, KeyError):
            tmp = self._doc.add_paragraph(style="List Bullet")
            tmp._element.getparent().remove(tmp._element)
        abstract_id = self._find_abstract_num_for_style("List Bullet")
        if abstract_id is None:
            for item in items:
                p = self._doc.add_paragraph(style="List Bullet")
                pf = p.paragraph_format
                pf.line_spacing = LINE_SPACING_BODY
                pf.space_after = Pt(0)
                run = p.add_run(_sanitize_prose(item))
                _set_run_font(run)
            return
        num_id = self._create_independent_num(abstract_id)
        for item in items:
            self._add_list_paragraph(item, num_id)

    def page_break(self):
        p = self._doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        self._just_broke_page = True

    def save(self, path: Optional[Union[str, Path]] = None) -> Path:
        """Сохранить .docx. Если path не задан — `<project>/docs/report.docx`.
        Относительный путь — от `<project>/docs/`. Абсолютный — без изменений.
        Создаёт parent-директории при необходимости. Возвращает абсолютный Path.

        После записи файл валидируется автоматически через `validate.py`.
        Если найдены нарушения уровня FAIL — поднимается `GostValidationError`.
        Тёплые предупреждения (heuristic) печатаются в stderr и не блокируют.
        Если validate.py недоступен (битый venv) — save() работает как раньше,
        отсутствие валидации лучше падения библиотеки.
        """
        if path is None:
            target = self._paths.out
        else:
            p = Path(path)
            target = p if p.is_absolute() else self._paths.docs / p
        target.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(target))
        target = target.resolve()

        try:
            import validate as _validate
        except ImportError:
            sys.stderr.write(
                "gost-report: validate.py недоступен, проверка пропущена.\n"
            )
            return target

        try:
            profile_name = type(self._profile).__name__
            for var_name, value in globals().items():
                if value is self._profile and var_name.endswith("_PROFILE"):
                    profile_name = var_name
                    break
            _validate.write_sentinel(target, profile_name=profile_name)
            violations = _validate.validate_docx(target)
        except Exception as e:
            sys.stderr.write(
                f"gost-report: валидатор упал ({e}); проверка пропущена.\n"
            )
            return target

        warns = [v for v in violations if v.tier == _validate.TIER_WARN]
        for v in warns:
            sys.stderr.write("gost-report warn: " + v.format().lstrip() + "\n")

        if _validate.has_failures(violations):
            raise GostValidationError(
                _validate.format_error(target, violations)
            )
        return target

    @property
    def doc(self) -> _Document:
        return self._doc

    @property
    def profile(self) -> UniversityProfile:
        return self._profile
