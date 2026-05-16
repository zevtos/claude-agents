#!/usr/bin/env python3
"""
extract_docx.py — Phase 4 extractor for DOCX files.

CLI:
    extract_docx.py <input_docx> <output_md>
                    [--doc-id doc-NNN]
                    [--source-rel relative/path/in/corpus.docx]

Effect:
    Writes <output_md> with YAML frontmatter + Markdown body. Stdout receives
    a single JSON line summarizing the result.

Pipeline:
    1. mammoth.convert_to_html(input) — semantic conversion preserving
       headings, lists, tables, inline images, footnotes.
    2. markdownify(html) — HTML → Markdown.
    3. python-docx scout for `inline_images`, `tables`, `paragraphs`,
       `has_equations`, `has_tracked_changes` (already gathered by scout
       but we re-derive here so the extractor is callable standalone).

Why HTML→Markdown instead of mammoth.convert_to_markdown:
    mammoth's built-in markdown writer drops tables. The HTML route keeps
    them, even if mammoth's HTML for tables is plain (no header markers
    beyond <th> — but markdownify renders them as Markdown pipe tables).
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from _common import (  # noqa: E402
    clean_whitespace,
    count_tokens,
    emit_failure,
    emit_success,
    sanitize_heading,
    sha256_of,
    today_iso,
    tool_version_string,
    validate_source_rel,
    write_md,
)


EXTRACTOR_NAME = "mammoth+markdownify"
# Below this body length on a non-empty docx, flag a warning.
MIN_BODY_CHARS = 50


def _try_import():
    try:
        import mammoth  # type: ignore
        from markdownify import markdownify as md_from_html  # type: ignore
        from docx import Document  # type: ignore
        return mammoth, md_from_html, Document
    except Exception as e:
        emit_failure(f"required libraries unavailable: {e}")
        sys.exit(1)


def _scout_docx(path: Path) -> dict:
    """Light re-scan to populate frontmatter — duplicates scout_corpus.py
    logic but keeps the extractor invocable without scout output."""
    from docx import Document  # type: ignore
    info = {
        "paragraphs": 0,
        "inline_images": 0,
        "has_tables": False,
        "has_equations": False,
        "has_tracked_changes": False,
    }
    try:
        doc = Document(str(path))
        info["paragraphs"] = len(doc.paragraphs)
        info["inline_images"] = len(doc.inline_shapes)
        info["has_tables"] = len(doc.tables) > 0
        xml = doc.element.xml
        info["has_equations"] = "<m:oMath" in xml or "<m:oMathPara" in xml
        info["has_tracked_changes"] = "w:ins" in xml or "w:del" in xml
    except Exception:
        pass
    return info


def extract(input_docx: Path) -> tuple[str, dict, list[str]]:
    """Returns (markdown_body, frontmatter_extras, warnings)."""
    mammoth, md_from_html, _Document = _try_import()
    warnings: list[str] = []

    # Convert DOCX → HTML. By default mammoth embeds images as base64 data
    # URIs — this catastrophically inflates output (~megabytes per file with
    # diagrams) and is useless for LLM consumption. Replace each image with
    # a compact placeholder that preserves semantics (caption + count).
    img_idx = {"n": 0}

    def _image_placeholder(image):
        img_idx["n"] += 1
        alt = (image.alt_text or "").strip()
        if alt:
            label = f"image {img_idx['n']}: {alt}"
        else:
            label = f"image {img_idx['n']}"
        # Returning {"src": ""} would emit <img src="">; mammoth lets us
        # return any attribute dict. Empty src + alt is cheapest.
        return {"src": "", "alt": label[:120]}

    image_handler = mammoth.images.img_element(_image_placeholder)
    try:
        with input_docx.open("rb") as fh:
            result = mammoth.convert_to_html(fh, convert_image=image_handler)
    except Exception as e:
        raise RuntimeError(f"mammoth conversion failed: {e}") from e

    html = result.value or ""
    # mammoth's messages are tuples of (type, message); promote warnings.
    for m in result.messages or []:
        try:
            mt = m.type if hasattr(m, "type") else m[0]
            mm = m.message if hasattr(m, "message") else m[1]
        except Exception:
            mt = "warning"
            mm = str(m)
        # mammoth reports lots of style-mapping notices that are noise here.
        if mt == "warning" and "style" not in str(mm).lower():
            warnings.append(f"mammoth: {mm}"[:200])

    # HTML → Markdown. heading_style=ATX renders "# heading" instead of underline.
    md = md_from_html(html, heading_style="ATX", bullets="-",
                      strip=["style", "script"])
    body = clean_whitespace(md)

    if len(body.strip()) < MIN_BODY_CHARS:
        warnings.append(f"extracted body is unusually short ({len(body)} chars)")

    extras = _scout_docx(input_docx)
    headings = []
    for line in body.split("\n"):
        if line.startswith("# ") and len(line) < 200:
            sanitized = sanitize_heading(line[2:])
            if sanitized:
                headings.append(sanitized)
            if len(headings) >= 10:
                break
    extras["headings"] = headings

    return body, extras, warnings


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract DOCX → Markdown via mammoth + markdownify.")
    ap.add_argument("input_docx")
    ap.add_argument("output_md")
    ap.add_argument("--doc-id", default="doc-000")
    ap.add_argument("--source-rel", default=None)
    args = ap.parse_args()

    in_path = Path(args.input_docx).expanduser().resolve()
    out_path = Path(args.output_md).expanduser().resolve()
    source_rel = args.source_rel or in_path.name
    try:
        source_rel = validate_source_rel(source_rel)
    except ValueError as e:
        emit_failure(f"invalid --source-rel: {e}")
        return 1

    if not in_path.is_file():
        emit_failure(f"input not found: {in_path}")
        return 1

    try:
        body, extras, warnings = extract(in_path)
    except Exception as e:
        emit_failure(f"extraction failed: {e}", extra={"input": str(in_path)})
        return 1

    fm = {
        "id": args.doc_id,
        "source": source_rel,
        "source_type": "docx",
        "source_sha256": sha256_of(in_path),
        "extraction_method": f"{EXTRACTOR_NAME}@{tool_version_string()}",
        "extraction_date": today_iso(),
        "paragraphs": extras.get("paragraphs", 0),
        "inline_images": extras.get("inline_images", 0),
        "has_tables": extras.get("has_tables", False),
        "has_equations": extras.get("has_equations", False),
        "has_tracked_changes": extras.get("has_tracked_changes", False),
        "headings": extras.get("headings", []),
        "tokens_estimated": count_tokens(body),
        "warnings": warnings,
    }
    write_md(out_path, fm, body)
    emit_success(out_path, body, warnings, extra={
        "inline_images": extras.get("inline_images", 0),
        "has_tables": extras.get("has_tables", False),
    })
    return 0


if __name__ == "__main__":
    sys.exit(main())
